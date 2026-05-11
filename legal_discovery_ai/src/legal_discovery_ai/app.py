from __future__ import annotations

import io
import json
import threading
import time
from datetime import datetime
from pathlib import Path
from typing import Any

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import streamlit as st

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None

from legal_discovery_ai.crew import PROJECT_ROOT, ensure_supported_python, run_crew


UPLOAD_DIR = PROJECT_ROOT / "data" / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
HISTORY_FILE = PROJECT_ROOT / "data" / "analysis_history.json"
APP_VERSION = "0.3.0"


def save_uploaded_pdf(uploaded_file: Any) -> Path:
    destination = UPLOAD_DIR / uploaded_file.name
    destination.write_bytes(uploaded_file.getbuffer())
    return destination


def normalize_result(result: object) -> str:
    if isinstance(result, str):
        return result
    if isinstance(result, dict):
        return json.dumps(result, indent=2, ensure_ascii=True)
    return str(result)


def build_user_friendly_error(exc: Exception) -> str:
    message = str(exc)
    lowered = message.lower()
    if "resource_exhausted" in lowered or "quota exceeded" in lowered:
        return (
            "Gemini quota is exhausted for this API key/project (free-tier limit reached "
            "or set to 0). Add billing in Google AI Studio/Cloud, switch to another "
            "provider key (Anthropic/OpenAI), or wait for quota reset before retrying."
        )
    if "not_found" in lowered and "gemini" in lowered and "model" in lowered:
        return (
            "Selected Gemini model is unavailable for this key/project. Set "
            "`GEMINI_MODEL` in `.env` to a model available in your account "
            "(for example `gemini-2.0-flash`)."
        )
    return message


def parse_result_payload(result: object) -> dict[str, Any]:
    if isinstance(result, dict):
        return result
    if isinstance(result, str):
        try:
            parsed = json.loads(result)
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError:
            return {"case_summary": result}
    return {"case_summary": str(result)}


def as_bullets(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item) for item in value if str(item).strip()]
    if isinstance(value, dict):
        return [f"{k}: {v}" for k, v in value.items()]
    text = str(value).strip()
    return [text] if text else []


def load_history() -> list[dict[str, str]]:
    if not HISTORY_FILE.exists():
        return []
    try:
        return json.loads(HISTORY_FILE.read_text(encoding="utf-8"))
    except Exception:
        return []


def save_history_entry(filename: str) -> None:
    history = load_history()
    history.insert(
        0,
        {
            "filename": filename,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
        },
    )
    HISTORY_FILE.parent.mkdir(parents=True, exist_ok=True)
    HISTORY_FILE.write_text(
        json.dumps(history[:20], indent=2, ensure_ascii=True), encoding="utf-8"
    )


def inject_custom_css() -> None:
    st.markdown(
        """
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
        html, body, [class*="css"] {
          font-family: 'Inter', sans-serif;
        }
        .stApp {
          background: radial-gradient(1200px 800px at 20% -10%, #1a2a52 0%, #0b1222 45%, #070d19 100%);
          color: #e6ecff;
        }
        .hero-card {
          border: 1px solid rgba(110, 146, 255, 0.28);
          border-radius: 18px;
          padding: 18px 20px;
          background: linear-gradient(135deg, rgba(27,38,72,0.72), rgba(14,20,38,0.78));
          box-shadow: 0 10px 28px rgba(4, 7, 18, 0.35);
          margin-bottom: 12px;
        }
        .soft-card {
          border: 1px solid rgba(115, 136, 209, 0.22);
          border-radius: 14px;
          padding: 14px;
          background: rgba(16, 24, 46, 0.76);
        }
        .risk-high {
          background: rgba(216, 58, 76, 0.22);
          border: 1px solid rgba(240, 93, 112, 0.55);
          border-radius: 10px;
          padding: 8px 10px;
          margin-bottom: 8px;
        }
        .risk-medium {
          background: rgba(212, 139, 53, 0.18);
          border: 1px solid rgba(255, 173, 74, 0.5);
          border-radius: 10px;
          padding: 8px 10px;
          margin-bottom: 8px;
        }
        .timeline-item {
          border-left: 2px solid rgba(111, 153, 255, 0.65);
          padding: 6px 0 6px 12px;
          margin-left: 6px;
          margin-bottom: 8px;
        }
        div[data-testid="stFileUploader"] {
          border: 2px dashed rgba(117, 154, 255, 0.55);
          border-radius: 16px;
          padding: 18px;
          background: rgba(13, 21, 41, 0.68);
        }
        .app-footer {
          margin-top: 30px;
          text-align: center;
          color: #8fa2d7;
          font-size: 0.9rem;
          border-top: 1px solid rgba(126, 145, 207, 0.22);
          padding-top: 12px;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def render_copy_button(text: str, key: str, label: str) -> None:
    if st.button(label, key=f"copy_btn_{key}"):
        st.toast("Copy ready. Use Ctrl+C from the raw JSON section.")


def render_bullet_section(title: str, value: Any) -> None:
    st.markdown(f"#### {title}")
    items = as_bullets(value)
    if not items:
        st.caption("No items identified.")
        return
    for item in items:
        st.markdown(f"- {item}")


def build_report_markdown(result_payload: dict[str, Any]) -> str:
    def section(title: str, value: Any) -> str:
        items = as_bullets(value)
        if not items:
            return f"## {title}\n- No items provided.\n"
        lines = "\n".join(f"- {item}" for item in items)
        return f"## {title}\n{lines}\n"

    parts = [
        "# Legal Discovery Case Brief",
        "",
        "## Case Summary",
        str(result_payload.get("case_summary", "No case summary available.")),
        "",
        section("Timeline", result_payload.get("timeline")),
        section("Parties", result_payload.get("parties")),
        section("Key Issues", result_payload.get("key_issues")),
        section("Critical Risks", result_payload.get("critical_risks")),
        section("Missing Elements", result_payload.get("missing_elements")),
        section("Next Actions", result_payload.get("next_actions")),
    ]
    return "\n".join(parts)


def markdown_to_pdf_bytes(markdown_text: str) -> bytes:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    _, height = letter
    y = height - 50
    max_chars = 100

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            y -= 12
        else:
            clean = line.replace("## ", "").replace("# ", "")
            while len(clean) > max_chars:
                pdf.drawString(50, y, clean[:max_chars])
                clean = clean[max_chars:]
                y -= 14
                if y < 50:
                    pdf.showPage()
                    y = height - 50
            pdf.drawString(50, y, clean)
            y -= 14

        if y < 50:
            pdf.showPage()
            y = height - 50

    pdf.save()
    buffer.seek(0)
    return buffer.read()


def markdown_to_docx_bytes(markdown_text: str) -> bytes | None:
    if Document is None:
        return None
    doc = Document()
    for line in markdown_text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line.replace("# ", "", 1), level=1)
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", "", 1), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line.replace("- ", "", 1), style="List Bullet")
        else:
            doc.add_paragraph(line)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def render_status_panel(
    progress_slot: Any, panel_slot: Any, statuses: dict[str, str], pct: int
) -> None:
    progress_slot.progress(pct / 100, text=f"Estimated progress: {pct}%")
    with panel_slot.container():
        st.markdown("### Agent Execution Status")
        for agent_name, status in statuses.items():
            icon = "[...]" if status == "Running" else "[x]" if status == "Completed" else "[ ]"
            st.write(f"{icon} {agent_name}: {status}")


def main() -> None:
    try:
        ensure_supported_python()
    except RuntimeError as exc:
        st.error(str(exc))
        st.stop()

    st.set_page_config(page_title="Legal Discovery AI", layout="wide")
    inject_custom_css()

    st.markdown(
        """
        <div class="hero-card">
          <h1 style="margin:0;">Legal Discovery AI</h1>
          <p style="margin:6px 0 0;color:#b5c7ff;">
            AI-Powered Discovery for DC Solo Attorneys
          </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.sidebar:
        st.header("Configuration")
        supplemental_text = st.text_area(
            "Supplemental instructions",
            value="",
            help="Optional direction for legal strategy, jurisdiction focus, or issue priorities.",
        )
        st.markdown("---")
        st.subheader("Analysis History")
        history = load_history()
        if history:
            for entry in history[:8]:
                st.caption(f"{entry['timestamp']} - {entry['filename']}")
        else:
            st.caption("No analyses yet.")
        st.markdown("---")
        if st.button("Add to RAG Knowledge Base", use_container_width=True):
            st.toast("Place PDF files in data/past_cases and rerun analysis.")
        if st.button("View Past Cases", use_container_width=True):
            st.toast("Open the data/past_cases folder in your workspace.")
        run_btn = st.button("Run Analysis", type="primary")

    st.markdown(
        '<div class="soft-card"><strong>Upload Case Document</strong><br/>'
        'Drop PDF here or click to upload</div>',
        unsafe_allow_html=True,
    )
    uploaded_pdf = st.file_uploader(
        "Drop PDF here or click to upload",
        type=["pdf"],
        label_visibility="collapsed",
    )
    if uploaded_pdf is not None:
        size_kb = uploaded_pdf.size / 1024
        st.markdown(
            f"""
            <div class="soft-card" style="margin-top:10px;">
                <strong>Document Ready</strong><br/>
                File: {uploaded_pdf.name}<br/>
                Size: {size_kb:.1f} KB
            </div>
            """,
            unsafe_allow_html=True,
        )

    if run_btn:
        if uploaded_pdf is None:
            st.error("Please upload a PDF before running analysis.")
            return

        document_path = save_uploaded_pdf(uploaded_pdf)
        save_history_entry(uploaded_pdf.name)
        st.info(f"Running crew on `{document_path}`...")

        statuses = {
            "Document Parser": "Queued",
            "Risk Scanner": "Queued",
            "Case Brief Writer": "Queued",
        }
        progress_slot = st.empty()
        status_panel_slot = st.empty()

        result_box: dict[str, Any] = {"value": None, "error": None}

        def worker() -> None:
            try:
                result_box["value"] = run_crew(
                    document_path=str(document_path),
                    document_text=supplemental_text or None,
                )
            except Exception as exc:  # pragma: no cover
                result_box["error"] = exc

        thread = threading.Thread(target=worker, daemon=True)
        thread.start()

        phase_order = ["Document Parser", "Risk Scanner", "Case Brief Writer"]
        phase_idx = 0
        start_time = time.time()

        while thread.is_alive():
            elapsed = time.time() - start_time
            estimated_phase = min(int(elapsed // 10), len(phase_order) - 1)
            phase_idx = max(phase_idx, estimated_phase)

            for idx, agent_name in enumerate(phase_order):
                if idx < phase_idx:
                    statuses[agent_name] = "Completed"
                elif idx == phase_idx:
                    statuses[agent_name] = "Running"
                else:
                    statuses[agent_name] = "Queued"

            pct = min(90, int((elapsed / 30) * 100))
            render_status_panel(
                progress_slot, status_panel_slot, statuses, max(5, pct)
            )
            time.sleep(0.5)

        thread.join()

        if result_box["error"] is not None:
            st.error(build_user_friendly_error(result_box["error"]))
            with st.expander("Technical details"):
                st.exception(result_box["error"])
            st.toast("Analysis failed. Review details and retry.")
            return

        for agent_name in phase_order:
            statuses[agent_name] = "Completed"
        render_status_panel(progress_slot, status_panel_slot, statuses, 100)

        result = result_box["value"]
        result_text = normalize_result(result)
        result_payload = parse_result_payload(result)

        st.toast("Analysis complete.")
        st.success("Analysis complete.")
        confidence_scores = result_payload.get("risk_scores", {})
        if isinstance(confidence_scores, dict) and confidence_scores:
            st.markdown("### Confidence Scores")
            cols = st.columns(min(4, len(confidence_scores)))
            for idx, (k, v) in enumerate(confidence_scores.items()):
                cols[idx % len(cols)].metric(str(k), str(v))

        facts_tab, risk_tab, brief_tab = st.tabs(
            ["Extracted Facts", "Risk Analysis", "Case Brief"]
        )

        with facts_tab:
            st.markdown("### Extracted Facts")
            render_bullet_section("Entities", result_payload.get("entities"))
            render_bullet_section("Parties", result_payload.get("parties"))
            render_bullet_section("Metadata", result_payload.get("metadata"))
            render_copy_button(result_text, "0", "Copy Facts")

        with risk_tab:
            st.markdown("### Risk Analysis")
            for risk in as_bullets(result_payload.get("critical_risks")):
                risk_lower = risk.lower()
                if any(x in risk_lower for x in ("high", "critical", "severe")):
                    st.markdown(
                        f'<div class="risk-high">{risk}</div>', unsafe_allow_html=True
                    )
                else:
                    st.markdown(
                        f'<div class="risk-medium">{risk}</div>', unsafe_allow_html=True
                    )
            render_bullet_section("Privilege Flags", result_payload.get("privilege_flags"))
            render_bullet_section("Regulatory Issues", result_payload.get("regulatory_issues"))
            render_copy_button(result_text, "1", "Copy Risks")

        with brief_tab:
            st.markdown("### Case Brief")
            st.markdown("#### Summary")
            st.write(result_payload.get("case_summary", "No case summary available."))
            st.markdown("#### Timeline")
            timeline_items = as_bullets(result_payload.get("timeline"))
            if timeline_items:
                for item in timeline_items:
                    st.markdown(
                        f'<div class="timeline-item">{item}</div>',
                        unsafe_allow_html=True,
                    )
            else:
                st.caption("No timeline extracted.")
            render_bullet_section("Key Issues", result_payload.get("key_issues"))
            render_bullet_section("Next Actions", result_payload.get("next_actions"))
            render_copy_button(result_text, "2", "Copy Brief")

        st.subheader("Export Options")
        report_markdown = build_report_markdown(result_payload)
        report_pdf = markdown_to_pdf_bytes(report_markdown)
        report_docx = markdown_to_docx_bytes(report_markdown)
        export_format = st.radio(
            "Choose export format",
            ["Markdown", "PDF", "Word (.docx)", "JSON"],
            horizontal=True,
        )
        if export_format == "Markdown":
            st.download_button(
                label="One-Click Export",
                data=report_markdown,
                file_name="case_brief_report.md",
                mime="text/markdown",
            )
        elif export_format == "PDF":
            st.download_button(
                label="One-Click Export",
                data=report_pdf,
                file_name="case_brief_report.pdf",
                mime="application/pdf",
            )
        elif export_format == "Word (.docx)":
            if report_docx is None:
                st.info("Word export unavailable. Install `python-docx` to enable it.")
            else:
                st.download_button(
                    label="One-Click Export",
                    data=report_docx,
                    file_name="case_brief_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
        else:
            st.download_button(
                label="One-Click Export",
                data=result_text,
                file_name="case_brief_output.json",
                mime="application/json",
            )

        with st.expander("Raw JSON Output"):
            st.code(result_text, language="json")
        st.download_button(
            label="Download output JSON",
            data=result_text,
            file_name="case_brief_output.json",
            mime="application/json",
        )

    st.markdown(
        f'<div class="app-footer">Built for DC Attorneys | Version {APP_VERSION}</div>',
        unsafe_allow_html=True,
    )


if __name__ == "__main__":
    main()
